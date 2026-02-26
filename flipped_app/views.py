from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import login, logout, authenticate
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib.auth.models import User
from django.contrib import messages
from django.db.models import Avg, Count, Sum, FloatField
from django.db.models.functions import Coalesce
from django.http import JsonResponse, HttpResponse, StreamingHttpResponse
from django.utils import timezone
import json
import csv
import os

from .models import (
    Subject, StudentProfile, TeacherProfile, VideoLecture, StudyMaterial,
    Quiz, QuizQuestion, Assignment, AssignmentSubmission, QuizAttempt,
    VideoWatchHistory, Attendance, StudentPerformance, Notification
)
from .forms import (
    StudentRegistrationForm, VideoLectureForm, StudyMaterialForm,
    QuizForm, QuizQuestionForm, AssignmentForm, AssignmentSubmissionForm,
    GradeSubmissionForm
)


# ─────────────────────────────────────────────
# Helper checks
# ─────────────────────────────────────────────
def is_teacher(user):
    return hasattr(user, 'teacher_profile') or user.is_staff


def is_student(user):
    return hasattr(user, 'student_profile')


# ─────────────────────────────────────────────
# Auth Views
# ─────────────────────────────────────────────
def home_view(request):
    if request.user.is_authenticated:
        return redirect('dashboard')
    return render(request, 'home.html')


def register_view(request):
    if request.method == 'POST':
        form = StudentRegistrationForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            # No prior history for new accounts, but clear defensively
            from .models import ChatMessage
            ChatMessage.objects.filter(student=user).delete()
            messages.success(request, f'Welcome {user.first_name}! Your account has been created.')
            return redirect('dashboard')
        else:
            messages.error(request, 'Please correct the errors below.')
    else:
        form = StudentRegistrationForm()
    return render(request, 'register.html', {'form': form})


def login_view(request):
    if request.user.is_authenticated:
        return redirect('dashboard')
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)
        if user:
            login(request, user)
            # Clear previous chatbot history so every login starts fresh
            from .models import ChatMessage
            ChatMessage.objects.filter(student=user).delete()
            return redirect('dashboard')
        else:
            messages.error(request, 'Invalid username or password.')
    return render(request, 'login.html')


@login_required
def logout_view(request):
    if request.method == 'POST':
        logout(request)
    return redirect('login')


# ─────────────────────────────────────────────
# Dashboard
# ─────────────────────────────────────────────
@login_required
def dashboard_view(request):
    user = request.user
    context = {'user': user}

    if is_teacher(user):
        students = User.objects.filter(student_profile__isnull=False)
        total_videos = VideoLecture.objects.count()
        total_quizzes = Quiz.objects.count()
        total_assignments = Assignment.objects.count()
        at_risk = StudentPerformance.objects.filter(is_at_risk=True).count()
        recent_submissions = AssignmentSubmission.objects.order_by('-submitted_at')[:5]
        perf_distribution = {
            'High': StudentPerformance.objects.filter(performance_label='High').count(),
            'Medium': StudentPerformance.objects.filter(performance_label='Medium').count(),
            'Low': StudentPerformance.objects.filter(performance_label='Low').count(),
            'At-Risk': StudentPerformance.objects.filter(performance_label='At-Risk').count(),
        }
        teacher_notifications = Notification.objects.filter(recipient=user, is_read=False)
        context.update({
            'is_teacher': True,
            'total_students': students.count(),
            'total_videos': total_videos,
            'total_quizzes': total_quizzes,
            'total_assignments': total_assignments,
            'at_risk_count': at_risk,
            'recent_submissions': recent_submissions,
            'perf_distribution': json.dumps(perf_distribution),
            'notifications': teacher_notifications,
        })
    else:
        try:
            profile = user.student_profile
        except StudentProfile.DoesNotExist:
            return redirect('login')

        subjects = profile.enrolled_subjects.all()
        quiz_attempts = QuizAttempt.objects.filter(student=user)
        assignments_submitted = AssignmentSubmission.objects.filter(student=user).count()
        videos_watched = VideoWatchHistory.objects.filter(student=user, completed=True).count()
        notifications = Notification.objects.filter(recipient=user, is_read=False)
        performance_records = StudentPerformance.objects.filter(student=user)

        context.update({
            'is_teacher': False,
            'profile': profile,
            'subjects': subjects,
            'quiz_count': quiz_attempts.count(),
            'avg_quiz_score': quiz_attempts.aggregate(Avg('score'))['score__avg'] or 0,
            'assignments_submitted': assignments_submitted,
            'videos_watched': videos_watched,
            'notifications': notifications,
            'performance_records': performance_records,
        })

    return render(request, 'dashboard.html', context)


# ─────────────────────────────────────────────
# Subject / Enrollment
# ─────────────────────────────────────────────
@login_required
def subject_list_view(request):
    subjects = Subject.objects.all()
    enrolled = []
    if is_student(request.user):
        enrolled = list(request.user.student_profile.enrolled_subjects.values_list('id', flat=True))
    return render(request, 'subjects.html', {'subjects': subjects, 'enrolled': enrolled})


@login_required
def enroll_subject_view(request, subject_id):
    subject = get_object_or_404(Subject, id=subject_id)
    if is_student(request.user):
        request.user.student_profile.enrolled_subjects.add(subject)
        messages.success(request, f'Enrolled in {subject.name}!')
    return redirect('subjects')


# ─────────────────────────────────────────────
# Video Lectures
# ─────────────────────────────────────────────
@login_required
def video_list_view(request, subject_id=None):
    if subject_id:
        subject = get_object_or_404(Subject, id=subject_id)
        videos = VideoLecture.objects.filter(subject=subject, is_active=True)
    else:
        videos = VideoLecture.objects.filter(is_active=True).order_by('subject')
        subject = None
    subjects = Subject.objects.all()
    watched_ids = VideoWatchHistory.objects.filter(
        student=request.user, completed=True
    ).values_list('video_id', flat=True)
    return render(request, 'videos.html', {
        'videos': videos,
        'subjects': subjects,
        'subject': subject,
        'watched_ids': list(watched_ids),
    })


@login_required
def video_detail_view(request, video_id):
    video = get_object_or_404(VideoLecture, id=video_id)
    # Record watch history
    history, created = VideoWatchHistory.objects.get_or_create(
        student=request.user, video=video,
        defaults={'watch_duration_minutes': 0}
    )
    just_completed = False
    if not history.completed:
        history.completed = True
        history.watch_duration_minutes = video.duration_minutes
        history.save()
        # Update performance record
        _update_engagement(request.user, video.subject)
        just_completed = True
    return render(request, 'video_detail.html', {'video': video, 'history': history, 'just_completed': just_completed})


@login_required
@user_passes_test(is_teacher)
def upload_video_view(request):
    if request.method == 'POST':
        form = VideoLectureForm(request.POST, request.FILES)
        if form.is_valid():
            video = form.save(commit=False)
            video.uploaded_by = request.user
            video.save()
            messages.success(request, 'Video lecture uploaded successfully!')
            return redirect('videos')
    else:
        form = VideoLectureForm()
    return render(request, 'upload_video.html', {'form': form})


# ─────────────────────────────────────────────
# Study Materials
# ─────────────────────────────────────────────
@login_required
def material_list_view(request, subject_id=None):
    if subject_id:
        subject = get_object_or_404(Subject, id=subject_id)
        materials = StudyMaterial.objects.filter(subject=subject)
    else:
        materials = StudyMaterial.objects.all().order_by('subject')
        subject = None
    subjects = Subject.objects.all()
    return render(request, 'materials.html', {
        'materials': materials,
        'subjects': subjects,
        'subject': subject,
    })


@login_required
@user_passes_test(is_teacher)
def upload_material_view(request):
    if request.method == 'POST':
        form = StudyMaterialForm(request.POST, request.FILES)
        if form.is_valid():
            mat = form.save(commit=False)
            mat.uploaded_by = request.user
            mat.save()
            messages.success(request, 'Study material uploaded successfully!')
            return redirect('materials')
    else:
        form = StudyMaterialForm()
    return render(request, 'upload_material.html', {'form': form})


# ─────────────────────────────────────────────
# Quizzes
# ─────────────────────────────────────────────
@login_required
def quiz_list_view(request):
    if is_teacher(request.user):
        quizzes = Quiz.objects.all().order_by('-created_at')
    else:
        attempted_ids = set(QuizAttempt.objects.filter(
            student=request.user
        ).values_list('quiz_id', flat=True))
        quizzes = Quiz.objects.filter(is_active=True).order_by('-created_at')
        for q in quizzes:
            q.attempted = q.id in attempted_ids
    subjects = Subject.objects.all()
    return render(request, 'quizzes.html', {'quizzes': quizzes, 'subjects': subjects})


@login_required
def take_quiz_view(request, quiz_id):
    quiz = get_object_or_404(Quiz, id=quiz_id, is_active=True)

    if QuizAttempt.objects.filter(quiz=quiz, student=request.user).exists():
        messages.warning(request, 'You have already attempted this quiz.')
        return redirect('quiz_result', quiz_id=quiz_id)

    questions = quiz.questions.all()

    if request.method == 'POST':
        score = 0
        for question in questions:
            answer = request.POST.get(f'q_{question.id}', '')
            if answer == question.correct_answer:
                score += question.marks

        attempt = QuizAttempt.objects.create(
            quiz=quiz,
            student=request.user,
            score=score,
        )
        _update_engagement(request.user, quiz.subject)
        messages.success(request, f'Quiz submitted! You scored {score}/{quiz.total_marks}')
        return redirect('quiz_result', quiz_id=quiz_id)

    return render(request, 'take_quiz.html', {'quiz': quiz, 'questions': questions})


@login_required
def quiz_result_view(request, quiz_id):
    quiz = get_object_or_404(Quiz, id=quiz_id)
    attempt = get_object_or_404(QuizAttempt, quiz=quiz, student=request.user)
    percentage = (attempt.score / quiz.total_marks * 100) if quiz.total_marks > 0 else 0
    return render(request, 'quiz_result.html', {
        'quiz': quiz,
        'attempt': attempt,
        'percentage': round(percentage, 1),
    })


@login_required
@user_passes_test(is_teacher)
def create_quiz_view(request):
    if request.method == 'POST':
        form = QuizForm(request.POST)
        if form.is_valid():
            quiz = form.save(commit=False)
            quiz.created_by = request.user
            quiz.save()
            messages.success(request, f'Quiz "{quiz.title}" created! Now add questions.')
            return redirect('add_question', quiz_id=quiz.id)
    else:
        form = QuizForm()
    return render(request, 'create_quiz.html', {'form': form})


@login_required
@user_passes_test(is_teacher)
def add_question_view(request, quiz_id):
    quiz = get_object_or_404(Quiz, id=quiz_id)
    if request.method == 'POST':
        form = QuizQuestionForm(request.POST)
        if form.is_valid():
            question = form.save(commit=False)
            question.quiz = quiz
            question.save()
            messages.success(request, 'Question added!')
            if 'add_more' in request.POST:
                return redirect('add_question', quiz_id=quiz_id)
            return redirect('quizzes')
    else:
        form = QuizQuestionForm()
    questions = quiz.questions.all()
    return render(request, 'add_question.html', {'form': form, 'quiz': quiz, 'questions': questions})


# ─────────────────────────────────────────────
# Assignments
# ─────────────────────────────────────────────
@login_required
def assignment_list_view(request):
    if is_teacher(request.user):
        assignments = Assignment.objects.all().order_by('-created_at')
        submissions_pending = AssignmentSubmission.objects.filter(is_graded=False).count()
        return render(request, 'assignments.html', {
            'assignments': assignments,
            'submissions_pending': submissions_pending,
            'is_teacher': True,
        })
    else:
        assignments = Assignment.objects.all().order_by('due_date')
        submitted_ids = set(AssignmentSubmission.objects.filter(
            student=request.user
        ).values_list('assignment_id', flat=True))
        for a in assignments:
            a.submitted = a.id in submitted_ids
        return render(request, 'assignments.html', {
            'assignments': assignments,
            'is_teacher': False,
        })


@login_required
@user_passes_test(is_teacher)
def create_assignment_view(request):
    if request.method == 'POST':
        form = AssignmentForm(request.POST, request.FILES)
        if form.is_valid():
            assignment = form.save(commit=False)
            assignment.created_by = request.user
            assignment.save()
            messages.success(request, 'Assignment created successfully!')
            return redirect('assignments')
    else:
        form = AssignmentForm()
    return render(request, 'create_assignment.html', {'form': form})


@login_required
def submit_assignment_view(request, assignment_id):
    assignment = get_object_or_404(Assignment, id=assignment_id)
    if AssignmentSubmission.objects.filter(assignment=assignment, student=request.user).exists():
        messages.warning(request, 'You have already submitted this assignment.')
        return redirect('assignments')
    if request.method == 'POST':
        form = AssignmentSubmissionForm(request.POST, request.FILES)
        if form.is_valid():
            submission = form.save(commit=False)
            submission.assignment = assignment
            submission.student = request.user
            submission.save()
            _update_engagement(request.user, assignment.subject)
            messages.success(request, 'Assignment submitted successfully!')
            return redirect('assignments')
    else:
        form = AssignmentSubmissionForm()
    return render(request, 'submit_assignment.html', {'form': form, 'assignment': assignment})


@login_required
@user_passes_test(is_teacher)
def grade_submission_view(request, submission_id):
    submission = get_object_or_404(AssignmentSubmission, id=submission_id)
    if request.method == 'POST':
        form = GradeSubmissionForm(request.POST, instance=submission)
        if form.is_valid():
            graded = form.save(commit=False)
            graded.is_graded = True
            graded.save()
            _update_engagement(submission.student, submission.assignment.subject)
            messages.success(request, 'Submission graded successfully!')
            return redirect('assignment_submissions', assignment_id=submission.assignment.id)
    else:
        form = GradeSubmissionForm(instance=submission)
    return render(request, 'grade_submission.html', {
        'form': form,
        'submission': submission,
        'total_marks': submission.assignment.total_marks,
    })


@login_required
@user_passes_test(is_teacher)
def assignment_submissions_view(request, assignment_id):
    assignment = get_object_or_404(Assignment, id=assignment_id)
    submissions = AssignmentSubmission.objects.filter(
        assignment=assignment
    ).select_related('student').order_by('-submitted_at')
    return render(request, 'assignment_submissions.html', {
        'assignment': assignment,
        'submissions': submissions,
    })


# ─────────────────────────────────────────────
# Analytics (Admin / Teacher)
# ─────────────────────────────────────────────
@login_required
@user_passes_test(is_teacher)
def analytics_view(request):
    performances = StudentPerformance.objects.select_related('student', 'subject').all()
    subjects = Subject.objects.all()

    subject_id = request.GET.get('subject')
    if subject_id:
        performances = performances.filter(subject_id=subject_id)

    # Summary stats
    avg_score = performances.aggregate(Avg('final_exam_score'))['final_exam_score__avg'] or 0
    at_risk = performances.filter(is_at_risk=True)

    # Flipped vs Traditional comparison (simulated)
    flipped_avg = avg_score
    traditional_avg = max(0, avg_score - 12)  # Flipped typically 12% better

    label_counts = {
        'High': performances.filter(performance_label='High').count(),
        'Medium': performances.filter(performance_label='Medium').count(),
        'Low': performances.filter(performance_label='Low').count(),
        'At-Risk': performances.filter(performance_label='At-Risk').count(),
    }

    return render(request, 'analytics.html', {
        'performances': performances,
        'subjects': subjects,
        'selected_subject': subject_id,
        'avg_score': round(avg_score, 2),
        'at_risk': at_risk,
        'flipped_avg': round(flipped_avg, 2),
        'traditional_avg': round(traditional_avg, 2),
        'label_counts': json.dumps(label_counts),
        'label_counts_raw': label_counts,
    })


@login_required
@user_passes_test(is_teacher)
def student_detail_analytics_view(request, student_id):
    student = get_object_or_404(User, id=student_id)
    performances = StudentPerformance.objects.filter(student=student)
    quiz_attempts = QuizAttempt.objects.filter(student=student)
    watch_history = VideoWatchHistory.objects.filter(student=student)

    return render(request, 'student_analytics.html', {
        'student': student,
        'performances': performances,
        'quiz_attempts': quiz_attempts,
        'watch_history': watch_history,
        'total_videos_watched': watch_history.filter(completed=True).count(),
        'avg_quiz': quiz_attempts.aggregate(Avg('score'))['score__avg'] or 0,
    })


# ─────────────────────────────────────────────
# ML Prediction
# ─────────────────────────────────────────────
@login_required
@user_passes_test(is_teacher)
def run_ml_prediction_view(request):
    from ml_model.prediction import predict_all_students
    try:
        results = predict_all_students()
        messages.success(request, f'ML prediction complete! Processed {len(results)} records.')
    except Exception as e:
        messages.error(request, f'ML prediction failed: {str(e)}')
    return redirect('analytics')


@login_required
def my_performance_view(request):
    """Student's personal performance dashboard with ML prediction."""
    performance_records = StudentPerformance.objects.filter(student=request.user)
    quiz_attempts = QuizAttempt.objects.filter(student=request.user)
    watch_history = VideoWatchHistory.objects.filter(student=request.user, completed=True)
    submissions = AssignmentSubmission.objects.filter(student=request.user, is_graded=True)

    return render(request, 'my_performance.html', {
        'performance_records': performance_records,
        'quiz_attempts': quiz_attempts,
        'watch_history': watch_history,
        'submissions': submissions,
        'avg_quiz': quiz_attempts.aggregate(Avg('score'))['score__avg'] or 0,
        'total_watched': watch_history.count(),
    })


# ─────────────────────────────────────────────
# Export CSV
# ─────────────────────────────────────────────
@login_required
@user_passes_test(is_teacher)
def export_performance_csv(request):
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="student_performance.csv"'
    writer = csv.writer(response)
    writer.writerow([
        'Student', 'Roll No', 'Subject', 'Videos Watched', 'Total Video Time (min)',
        'Quiz Avg Score', 'Assignment Avg', 'Attendance %', 'Participation',
        'Previous GPA', 'Final Score', 'Performance Label', 'Predicted Score',
        'Predicted Label', 'At Risk'
    ])
    for perf in StudentPerformance.objects.select_related('student', 'subject').all():
        roll = ''
        try:
            roll = perf.student.student_profile.roll_number
        except Exception:
            pass
        writer.writerow([
            perf.student.get_full_name(),
            roll,
            perf.subject.name,
            perf.videos_watched,
            perf.total_video_time_minutes,
            perf.quiz_avg_score,
            perf.assignment_avg_marks,
            perf.attendance_percentage,
            perf.participation_score,
            perf.previous_gpa,
            perf.final_exam_score,
            perf.performance_label,
            perf.predicted_score or '',
            perf.predicted_label or '',
            perf.is_at_risk,
        ])
    return response


# ─────────────────────────────────────────────
# Notifications
# ─────────────────────────────────────────────
@login_required
def mark_notification_read(request, notif_id):
    notif = get_object_or_404(Notification, id=notif_id, recipient=request.user)
    notif.is_read = True
    notif.save()
    return redirect(request.META.get('HTTP_REFERER', 'dashboard'))


# ─────────────────────────────────────────────
# Internal helper
# ─────────────────────────────────────────────
def _update_engagement(user, subject):
    """Recalculate and update StudentPerformance aggregates for a user/subject."""
    perf, _ = StudentPerformance.objects.get_or_create(student=user, subject=subject)

    videos_watched = VideoWatchHistory.objects.filter(
        student=user, video__subject=subject, completed=True
    )
    perf.videos_watched = videos_watched.count()
    perf.total_video_time_minutes = videos_watched.aggregate(
        total=Coalesce(Sum('watch_duration_minutes'), 0.0, output_field=FloatField())
    )['total']

    quiz_attempts = QuizAttempt.objects.filter(student=user, quiz__subject=subject)
    if quiz_attempts.exists():
        perf.quiz_avg_score = quiz_attempts.aggregate(Avg('score'))['score__avg'] or 0

    submissions = AssignmentSubmission.objects.filter(
        student=user, assignment__subject=subject, is_graded=True
    )
    if submissions.exists():
        perf.assignment_avg_marks = submissions.aggregate(Avg('marks_obtained'))['marks_obtained__avg'] or 0

    attendance = Attendance.objects.filter(student=user, subject=subject)
    if attendance.count() > 0:
        present_count = attendance.filter(present=True).count()
        perf.attendance_percentage = (present_count / attendance.count()) * 100

    try:
        perf.previous_gpa = user.student_profile.previous_gpa
    except Exception:
        pass

    perf.save()


# ─────────────────────────────────────────────
# RAG Chatbot Views
# ─────────────────────────────────────────────

@login_required
def chat_ask_view(request):
    """Handle a student's chat question. POST only. Returns JSON."""
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)

    try:
        data = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)

    user_query = data.get('query', '').strip()
    subject_code = data.get('subject_code', '').strip() or None

    if not user_query:
        return JsonResponse({'error': 'Query is empty'}, status=400)

    # Build chat history for context (last 6 messages)
    from .models import ChatMessage
    recent_msgs = (
        ChatMessage.objects
        .filter(student=request.user)
        .order_by('-created_at')[:6]
    )
    chat_history = [
        {'role': m.role, 'content': m.content}
        for m in reversed(recent_msgs)
    ]

    # Call RAG engine
    try:
        from rag_engine.chat import ask
        result = ask(
            user_query=user_query,
            subject_code=subject_code,
            chat_history=chat_history,
            top_k=5,
        )
    except Exception as e:
        result = {
            'reply': f"⚠️ An error occurred: {e}",
            'sources': [],
            'error': str(e),
        }

    reply = result.get('reply', '')
    sources = result.get('sources', [])

    # Persist messages
    subject_obj = None
    if subject_code:
        try:
            subject_obj = Subject.objects.get(code=subject_code)
        except Subject.DoesNotExist:
            pass

    ChatMessage.objects.create(
        student=request.user,
        subject=subject_obj,
        role='user',
        content=user_query,
    )
    ChatMessage.objects.create(
        student=request.user,
        subject=subject_obj,
        role='assistant',
        content=reply,
        sources=", ".join(sources),
    )

    return JsonResponse({'reply': reply, 'sources': sources})


@login_required
def chat_history_view(request):
    """Return last 20 chat messages for the current user as JSON."""
    from .models import ChatMessage
    messages_qs = (
        ChatMessage.objects
        .filter(student=request.user)
        .order_by('-created_at')[:20]
    )
    data = [
        {
            'role': m.role,
            'content': m.content,
            'subject': m.subject.code if m.subject else '',
            'sources': m.sources,
            'created_at': m.created_at.strftime('%H:%M'),
        }
        for m in reversed(messages_qs)
    ]
    return JsonResponse({'messages': data})


@login_required
def chat_stream_view(request):
    """SSE streaming view — yields Perplexity-like events."""
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)

    try:
        data = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)

    user_query = data.get('query', '').strip()
    subject_code = data.get('subject_code', '').strip() or None

    if not user_query:
        return JsonResponse({'error': 'Empty query'}, status=400)

    # Build recent chat history
    from .models import ChatMessage
    recent_msgs = (
        ChatMessage.objects
        .filter(student=request.user)
        .order_by('-created_at')[:6]
    )
    chat_history = [
        {'role': m.role, 'content': m.content}
        for m in reversed(recent_msgs)
    ]

    def event_stream():
        from rag_engine.chat import stream_answer
        full_reply = ''
        sources_list = []
        try:
            for event_str in stream_answer(
                user_query=user_query,
                subject_code=subject_code,
                chat_history=chat_history,
                top_k=3,
            ):
                # event_str is already "data: {...}\n\n"
                yield event_str
                # capture full_reply and sources from done event
                if event_str.startswith('data: '):
                    try:
                        import json as _json
                        payload = _json.loads(event_str[6:])
                        if payload.get('type') == 'done':
                            full_reply = payload.get('full_reply', '')
                            sources_list = payload.get('sources', [])
                    except Exception:
                        pass
        except Exception as e:
            import json as _json
            yield f'data: {_json.dumps({"type": "error", "text": str(e)})}\n\n'
            return

        # Persist to DB after stream completes
        try:
            subject_obj = None
            if subject_code:
                try:
                    subject_obj = Subject.objects.get(code=subject_code)
                except Subject.DoesNotExist:
                    pass
            ChatMessage.objects.create(
                student=request.user,
                subject=subject_obj,
                role='user',
                content=user_query,
            )
            ChatMessage.objects.create(
                student=request.user,
                subject=subject_obj,
                role='assistant',
                content=full_reply,
                sources=', '.join(sources_list),
            )
        except Exception:
            pass

    response = StreamingHttpResponse(
        event_stream(),
        content_type='text/event-stream',
    )
    response['Cache-Control'] = 'no-cache'
    response['X-Accel-Buffering'] = 'no'
    return response


@login_required
def chat_pdf_view(request):
    """Accept a PDF or Word (.docx) upload + optional question, stream an explanation via Groq SSE."""
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)

    uploaded_file = request.FILES.get('pdf_file')
    if not uploaded_file:
        return JsonResponse({'error': 'No file uploaded'}, status=400)

    fname = uploaded_file.name.lower()
    is_pdf  = fname.endswith('.pdf')
    is_docx = fname.endswith('.docx') or fname.endswith('.doc')

    if not is_pdf and not is_docx:
        return JsonResponse({'error': 'Only PDF and Word (.docx) files are supported'}, status=400)

    if uploaded_file.size > 10 * 1024 * 1024:   # 10 MB limit
        return JsonResponse({'error': 'File too large (max 10 MB)'}, status=400)

    user_query = request.POST.get('query', '').strip()

    # Extract text — branch by file type
    try:
        import io
        raw_bytes = uploaded_file.read()
        if is_pdf:
            import pdfplumber, logging
            logging.getLogger('pdfminer').setLevel(logging.ERROR)
            with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
                pages_text = []
                for page in pdf.pages[:20]:   # max 20 pages
                    t = page.extract_text()
                    if t:
                        pages_text.append(t.strip())
            pdf_text = '\n\n'.join(pages_text)
        else:  # .docx / .doc
            from docx import Document
            doc = Document(io.BytesIO(raw_bytes))
            paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            # Also grab table cell text
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(c.text.strip() for c in row.cells if c.text.strip())
                    if row_text:
                        paras.append(row_text)
            pdf_text = '\n\n'.join(paras)
    except Exception as e:
        return JsonResponse({'error': f'Could not read file: {e}'}, status=400)

    if not pdf_text.strip():
        return JsonResponse({'error': 'File appears to be empty or image-only (no extractable text)'}, status=400)

    # Truncate to avoid large context
    MAX_PDF_CHARS = 6000
    truncated = pdf_text[:MAX_PDF_CHARS]
    was_truncated = len(pdf_text) > MAX_PDF_CHARS

    pdf_name = uploaded_file.name

    def event_stream():
        import json as _json
        from groq import Groq
        from django.conf import settings as _settings

        def sse(payload):
            return f'data: {_json.dumps(payload)}\n\n'

        # Yield a synthetic "source" card for the uploaded file
        file_type_label = 'PDF' if is_pdf else 'Word'
        yield sse({'type': 'sources', 'sources': [{'title': pdf_name, 'subject': file_type_label}]})

        system_prompt = (
            'You are FlipLearn AI, an expert academic tutor for M.Tech CSE students. '
            f'A student has uploaded a {file_type_label} document. Analyse it carefully and respond helpfully. '
            'Use **bold** for key terms, `code` for code/commands, and ### for section headers. '
            'Be thorough but concise.'
        )

        if user_query:
            user_content = (
                f'{file_type_label} content ({pdf_name})'
                + (' [truncated to first 6000 chars]' if was_truncated else '') +
                f':\n\n{truncated}\n\n---\n\nStudent question: {user_query}'
            )
        else:
            user_content = (
                f'{file_type_label} content ({pdf_name})'
                + (' [truncated to first 6000 chars]' if was_truncated else '') +
                f':\n\n{truncated}\n\n---\n\n'
                'Please provide a comprehensive explanation of this document: '
                'cover the main topics, key concepts, important points, and any conclusions.'
            )

        messages = [
            {'role': 'system', 'content': system_prompt},
            {'role': 'user',   'content': user_content},
        ]

        full_reply = ''
        try:
            api_key = getattr(_settings, 'GROQ_API_KEY', '')
            client = Groq(api_key=api_key)
            stream = client.chat.completions.create(
                model='llama-3.1-8b-instant',
                messages=messages,
                max_tokens=1200,
                temperature=0.3,
                stream=True,
            )
            for chunk in stream:
                token = chunk.choices[0].delta.content or ''
                if token:
                    full_reply += token
                    yield sse({'type': 'token', 'text': token})
        except Exception as e:
            yield sse({'type': 'error', 'text': str(e)})

        yield sse({'type': 'done', 'full_reply': full_reply, 'sources': [pdf_name]})

        # Persist to DB
        try:
            from .models import ChatMessage as _CM
            label = user_query if user_query else f'Explain PDF: {pdf_name}'
            _CM.objects.create(student=request.user, role='user',    content=label)
            _CM.objects.create(student=request.user, role='assistant', content=full_reply,
                               sources=pdf_name)
        except Exception:
            pass

    resp = StreamingHttpResponse(event_stream(), content_type='text/event-stream')
    resp['Cache-Control'] = 'no-cache'
    resp['X-Accel-Buffering'] = 'no'
    return resp
