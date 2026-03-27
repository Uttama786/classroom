from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import login, logout, authenticate
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib.auth.models import User
from django.contrib import messages
from django.db.models import Avg, Count, Sum, FloatField
from django.db.models.functions import Coalesce
from django.http import JsonResponse, HttpResponse, StreamingHttpResponse, FileResponse, Http404
from django.conf import settings
from django.utils import timezone
from django.utils.http import url_has_allowed_host_and_scheme
from pathlib import Path
import json
import csv
import os
import logging
import threading

from .models import (
    Subject, StudentProfile, TeacherProfile, VideoLecture, StudyMaterial,
    Quiz, QuizQuestion, Assignment, AssignmentSubmission, QuizAttempt,
    VideoWatchHistory, Attendance, StudentPerformance, Notification, ChatMessage
)
from .forms import (
    StudentRegistrationForm, VideoLectureForm, StudyMaterialForm,
    QuizForm, QuizQuestionForm, AssignmentForm, AssignmentSubmissionForm,
    GradeSubmissionForm
)

logger = logging.getLogger(__name__)


# In-process state for RAG rebuild jobs.
RAG_REBUILD_LOCK = threading.Lock()
RAG_REBUILD_STATE = {
    'running': False,
    'started_at': None,
    'finished_at': None,
    'last_error': None,
}


def _rag_status_payload():
    return {
        'running': RAG_REBUILD_STATE['running'],
        'started_at': RAG_REBUILD_STATE['started_at'].isoformat() if RAG_REBUILD_STATE['started_at'] else None,
        'finished_at': RAG_REBUILD_STATE['finished_at'].isoformat() if RAG_REBUILD_STATE['finished_at'] else None,
        'last_error': RAG_REBUILD_STATE['last_error'],
    }


def _run_rag_rebuild_job():
    try:
        from rag_engine.indexer import build_index
        build_index()
        # Invalidate cached FAISS index so next query loads the new index
        from rag_engine.retriever import invalidate_cache
        invalidate_cache()
        err = None
    except Exception as exc:
        logger.exception('rebuild_rag_view failed: %s', exc)
        err = 'Failed to rebuild RAG index. Check server logs.'

    with RAG_REBUILD_LOCK:
        RAG_REBUILD_STATE['running'] = False
        RAG_REBUILD_STATE['finished_at'] = timezone.now()
        RAG_REBUILD_STATE['last_error'] = err


# ─────────────────────────────────────────────
# Helper checks
# ─────────────────────────────────────────────
def is_teacher(user):
    return hasattr(user, 'teacher_profile') or user.is_staff


def is_student(user):
    return hasattr(user, 'student_profile')


# ─────────────────────────────────────────────
# Auth Views
# ─────────────────────────────────────────────
def home_view(request):
    if request.user.is_authenticated:
        return redirect('dashboard')
    from .models import StudentProfile, VideoLecture, Quiz, StudyMaterial, QuizAttempt, StudentPerformance
    stats = {
        'total_students': StudentProfile.objects.count(),
        'total_videos':   VideoLecture.objects.count(),
        'total_quizzes':  Quiz.objects.count(),
        'total_materials': StudyMaterial.objects.count(),
        'total_quiz_attempts': QuizAttempt.objects.count(),
        'at_risk_count': StudentPerformance.objects.filter(is_at_risk=True).count(),
        'total_subjects': Subject.objects.count(),
        'total_assignments': Assignment.objects.count(),
    }
    features_list = [
        {'icon': 'bi-play-circle-fill', 'bg': 'rgba(37,99,235,.1)',   'color': '#2563eb', 'title': 'Video Lectures',           'desc': 'Watch curated videos before class at your own pace, any time, any device.'},
        {'icon': 'bi-patch-question-fill','bg':'rgba(22,163,74,.1)',  'color': '#16a34a', 'title': 'Adaptive Quizzes',          'desc': 'Test understanding with auto-graded quizzes tailored per subject.'},
        {'icon': 'bi-robot',              'bg': 'rgba(124,58,237,.1)','color': '#7c3aed', 'title': 'AI Tutor (RAG)',            'desc': 'Ask any CSE question — get instant, grounded answers from your course materials.'},
        {'icon': 'bi-graph-up-arrow',     'bg': 'rgba(202,138,4,.1)', 'color': '#ca8a04', 'title': 'ML Performance Prediction', 'desc': 'RandomForest models predict your score and classify your performance category.'},
        {'icon': 'bi-shield-exclamation', 'bg': 'rgba(220,38,38,.1)', 'color': '#dc2626', 'title': 'At-Risk Early Warning',     'desc': 'Detect struggling students early so teachers can intervene before it\'s too late.'},
        {'icon': 'bi-bar-chart-line-fill','bg': 'rgba(14,165,233,.1)','color': '#0ea5e9', 'title': 'Teacher Analytics',         'desc': 'Dashboards and charts showing class-wide and individual student progress.'},
        {'icon': 'bi-journal-code',       'bg': 'rgba(234,88,12,.1)', 'color': '#ea580c', 'title': 'Assignments',               'desc': 'Submit and grade coding assignments with inline feedback from teachers.'},
        {'icon': 'bi-file-earmark-text',  'bg': 'rgba(15,118,110,.1)','color': '#0f766e', 'title': 'Study Materials',           'desc': 'Download lecture notes, PDFs, and reference documents anytime.'},
    ]
    return render(request, 'home.html', {'stats': stats, 'features_list': features_list})


def register_view(request):
    if request.method == 'POST':
        form = StudentRegistrationForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            # No prior history for new accounts, but clear defensively
            ChatMessage.objects.filter(student=user).delete()
            messages.success(request, f'Welcome {user.first_name}! Your account has been created.')
            return redirect('dashboard')
        else:
            messages.error(request, 'Please correct the errors below.')
    else:
        form = StudentRegistrationForm()
    return render(request, 'register.html', {'form': form})


def login_view(request):
    if request.user.is_authenticated:
        return redirect('dashboard')
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)
        if user:
            login(request, user)
            # Clear previous chatbot history so every login starts fresh
            ChatMessage.objects.filter(student=user).delete()
            return redirect('dashboard')
        else:
            messages.error(request, 'Invalid username or password.')
    return render(request, 'login.html')


@login_required
def logout_view(request):
    if request.method == 'POST':
        logout(request)
        return redirect('login')
    # GET request to /logout/ — redirect to dashboard instead of causing a bounce loop
    return redirect('dashboard')


# ─────────────────────────────────────────────
# Dashboard
# ─────────────────────────────────────────────
@login_required
def dashboard_view(request):
    user = request.user
    context = {'user': user}

    # ── Admin dashboard ──────────────────────────────────────
    if user.is_superuser:
        total_students  = User.objects.filter(student_profile__isnull=False).count()
        total_teachers  = User.objects.filter(teacher_profile__isnull=False).count()
        total_videos    = VideoLecture.objects.count()
        total_materials = StudyMaterial.objects.count()
        total_quizzes   = Quiz.objects.count()
        total_assignments = Assignment.objects.count()
        at_risk         = StudentPerformance.objects.filter(is_at_risk=True).count()
        pending_grades  = AssignmentSubmission.objects.filter(is_graded=False).count()
        recent_users    = User.objects.order_by('-date_joined')[:8]
        perf_distribution = {
            'High':    StudentPerformance.objects.filter(performance_label='High').count(),
            'Medium':  StudentPerformance.objects.filter(performance_label='Medium').count(),
            'Low':     StudentPerformance.objects.filter(performance_label='Low').count(),
            'At-Risk': StudentPerformance.objects.filter(performance_label='At-Risk').count(),
        }
        subject_stats = [
            {
                'name':        s.name,
                'code':        s.code if hasattr(s, 'code') else '',
                'videos':      VideoLecture.objects.filter(subject=s).count(),
                'materials':   StudyMaterial.objects.filter(subject=s).count(),
                'quizzes':     Quiz.objects.filter(subject=s).count(),
                'assignments': Assignment.objects.filter(subject=s).count(),
                'enrolled':    s.studentprofile_set.count() if hasattr(s, 'studentprofile_set') else 0,
            }
            for s in Subject.objects.all()
        ]
        context.update({
            'role': 'admin',
            'total_students':    total_students,
            'total_teachers':    total_teachers,
            'total_videos':      total_videos,
            'total_materials':   total_materials,
            'total_quizzes':     total_quizzes,
            'total_assignments': total_assignments,
            'at_risk_count':     at_risk,
            'pending_grades':    pending_grades,
            'recent_users':      recent_users,
            'perf_distribution':      json.dumps(perf_distribution),
            'subject_stats':          subject_stats,
            'subject_stats_json':     json.dumps(subject_stats),
        })
        return render(request, 'dashboard_admin.html', context)

    # ── Teacher dashboard ────────────────────────────────────
    if is_teacher(user):
        students = User.objects.filter(student_profile__isnull=False)
        total_videos    = VideoLecture.objects.count()
        total_quizzes   = Quiz.objects.count()
        total_assignments = Assignment.objects.count()
        at_risk         = StudentPerformance.objects.filter(is_at_risk=True).count()
        pending_grades  = AssignmentSubmission.objects.filter(is_graded=False).count()
        recent_submissions = AssignmentSubmission.objects.order_by('-submitted_at')[:8]
        quiz_attempts_recent = QuizAttempt.objects.order_by('-attempted_at')[:8]
        perf_distribution = {
            'High':    StudentPerformance.objects.filter(performance_label='High').count(),
            'Medium':  StudentPerformance.objects.filter(performance_label='Medium').count(),
            'Low':     StudentPerformance.objects.filter(performance_label='Low').count(),
            'At-Risk': StudentPerformance.objects.filter(performance_label='At-Risk').count(),
        }
        teacher_notifications = Notification.objects.filter(recipient=user, is_read=False)

        # Subject-wise breakdown of enrolled students
        subject_enrollment = [
            {'subject': s, 'count': s.studentprofile_set.count()}
            for s in Subject.objects.all()
        ]

        context.update({
            'role': 'teacher',
            'is_teacher': True,
            'total_students':    students.count(),
            'total_videos':      total_videos,
            'total_quizzes':     total_quizzes,
            'total_assignments': total_assignments,
            'at_risk_count':     at_risk,
            'pending_grades':    pending_grades,
            'recent_submissions':    recent_submissions,
            'quiz_attempts_recent':  quiz_attempts_recent,
            'perf_distribution':     json.dumps(perf_distribution),
            'notifications':         teacher_notifications,
            'subject_enrollment':    subject_enrollment,
        })
        return render(request, 'dashboard_teacher.html', context)

    # ── Student dashboard ────────────────────────────────────
    try:
        profile = user.student_profile
    except StudentProfile.DoesNotExist:
        return redirect('login')

    subjects        = profile.enrolled_subjects.all()
    quiz_attempts   = QuizAttempt.objects.filter(student=user)
    assignments_submitted = AssignmentSubmission.objects.filter(student=user).count()
    videos_watched  = VideoWatchHistory.objects.filter(student=user, completed=True).count()
    notifications   = Notification.objects.filter(recipient=user, is_read=False)
    performance_records = StudentPerformance.objects.filter(student=user)
    recent_quiz_attempts = QuizAttempt.objects.filter(student=user).order_by('-attempted_at')[:5]
    recent_submissions   = AssignmentSubmission.objects.filter(student=user).order_by('-submitted_at')[:5]
    available_quizzes    = Quiz.objects.filter(is_active=True).exclude(
        attempts__student=user
    )[:4]

    context.update({
        'role': 'student',
        'is_teacher': False,
        'profile':       profile,
        'subjects':      subjects,
        'quiz_count':    quiz_attempts.count(),
        'avg_quiz_score': quiz_attempts.aggregate(Avg('score'))['score__avg'] or 0,
        'assignments_submitted':   assignments_submitted,
        'videos_watched':          videos_watched,
        'notifications':           notifications,
        'performance_records':     performance_records,
        'recent_quiz_attempts':    recent_quiz_attempts,
        'recent_submissions':      recent_submissions,
        'available_quizzes':       available_quizzes,
    })
    return render(request, 'dashboard_student.html', context)


# ─────────────────────────────────────────────
# Subject / Enrollment
# ─────────────────────────────────────────────
@login_required
def subject_list_view(request):
    subjects = Subject.objects.all()
    enrolled = []
    if is_student(request.user):
        enrolled = list(request.user.student_profile.enrolled_subjects.values_list('id', flat=True))
    return render(request, 'subjects.html', {'subjects': subjects, 'enrolled': enrolled})


@login_required
def enroll_subject_view(request, subject_id):
    if request.method != 'POST':
        return redirect('subjects')
    subject = get_object_or_404(Subject, id=subject_id)
    if is_student(request.user):
        request.user.student_profile.enrolled_subjects.add(subject)
        messages.success(request, f'Enrolled in {subject.name}!')
    else:
        messages.error(request, 'Only students can enrol in subjects.')
    return redirect('subjects')


# ─────────────────────────────────────────────
# Video Lectures
# ─────────────────────────────────────────────
@login_required
def video_list_view(request, subject_id=None):
    if subject_id:
        subject = get_object_or_404(Subject, id=subject_id)
        videos = VideoLecture.objects.filter(subject=subject, is_active=True)
    else:
        videos = VideoLecture.objects.filter(is_active=True).order_by('subject')
        subject = None
    subjects = Subject.objects.all()
    watched_ids = VideoWatchHistory.objects.filter(
        student=request.user, completed=True
    ).values_list('video_id', flat=True)
    return render(request, 'videos.html', {
        'videos': videos,
        'subjects': subjects,
        'subject': subject,
        'watched_ids': list(watched_ids),
    })


@login_required
def video_detail_view(request, video_id):
    video = get_object_or_404(VideoLecture, id=video_id)
    # Create a watch history entry but do NOT auto-complete —
    # the student must explicitly mark the video as watched.
    history, _ = VideoWatchHistory.objects.get_or_create(
        student=request.user, video=video,
        defaults={'watch_duration_minutes': 0, 'completed': False}
    )
    return render(request, 'video_detail.html', {'video': video, 'history': history})


@login_required
def mark_video_watched_view(request, video_id):
    """AJAX/POST endpoint: student explicitly marks a video as watched."""
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    if not is_student(request.user):
        return JsonResponse({'error': 'Only students can mark videos as watched'}, status=403)
    video = get_object_or_404(VideoLecture, id=video_id)
    history, _ = VideoWatchHistory.objects.get_or_create(
        student=request.user, video=video,
        defaults={'watch_duration_minutes': 0}
    )
    already_completed = history.completed
    if not history.completed:
        history.completed = True
        history.watch_duration_minutes = video.duration_minutes
        history.save()
        _update_engagement(request.user, video.subject)
    return JsonResponse({'status': 'ok', 'already_completed': already_completed})


@login_required
@user_passes_test(is_teacher)
def upload_video_view(request):
    if request.method == 'POST':
        form = VideoLectureForm(request.POST, request.FILES)
        if form.is_valid():
            video = form.save(commit=False)
            video.uploaded_by = request.user
            video.save()
            messages.success(request, 'Video lecture uploaded successfully!')
            return redirect('videos')
    else:
        form = VideoLectureForm()
    return render(request, 'upload_video.html', {'form': form})


# ─────────────────────────────────────────────
# Study Materials
# ─────────────────────────────────────────────
@login_required
def material_list_view(request, subject_id=None):
    if subject_id:
        subject = get_object_or_404(Subject, id=subject_id)
        materials = StudyMaterial.objects.filter(subject=subject)
    else:
        materials = StudyMaterial.objects.all().order_by('subject')
        subject = None
    subjects = Subject.objects.all()
    return render(request, 'materials.html', {
        'materials': materials,
        'subjects': subjects,
        'subject': subject,
    })


@login_required
def download_material_view(request, material_id):
    """Serve study materials from local MEDIA_ROOT with a stable app URL."""
    material = get_object_or_404(StudyMaterial, id=material_id)
    if not material.file or not material.file.name:
        raise Http404("Material file not found")

    filename = Path(material.file.name).name
    rel_path = Path("materials") / filename
    local_path = Path(settings.MEDIA_ROOT) / rel_path

    if not local_path.exists():
        raise Http404("Material file is unavailable on this server")

    # ── Track material download in StudentPerformance ─────────────────────
    if is_student(request.user):
        try:
            from django.db.models import F
            perf, _ = StudentPerformance.objects.get_or_create(
                student=request.user,
                subject=material.subject,
            )
            StudentPerformance.objects.filter(pk=perf.pk).update(
                materials_downloaded=F('materials_downloaded') + 1
            )
        except Exception:
            pass  # Never let tracking break the file download

    return FileResponse(
        local_path.open("rb"),
        as_attachment=True,
        filename=filename,
    )


@login_required
@user_passes_test(is_teacher)
def upload_material_view(request):
    if request.method == 'POST':
        form = StudyMaterialForm(request.POST, request.FILES)
        if form.is_valid():
            mat = form.save(commit=False)
            mat.uploaded_by = request.user
            mat.save()
            messages.success(request, 'Study material uploaded successfully!')
            return redirect('materials')
    else:
        form = StudyMaterialForm()
    return render(request, 'upload_material.html', {'form': form})


# ─────────────────────────────────────────────
# Quizzes
# ─────────────────────────────────────────────
@login_required
def quiz_list_view(request):
    if is_teacher(request.user):
        quizzes = Quiz.objects.all().order_by('-created_at')
    else:
        attempted_ids = set(QuizAttempt.objects.filter(
            student=request.user
        ).values_list('quiz_id', flat=True))
        quizzes = Quiz.objects.filter(is_active=True).order_by('-created_at')
        for q in quizzes:
            q.attempted = q.id in attempted_ids
    subjects = Subject.objects.all()
    return render(request, 'quizzes.html', {'quizzes': quizzes, 'subjects': subjects})


@login_required
def take_quiz_view(request, quiz_id):
    if not is_student(request.user):
        messages.error(request, 'Only students can attempt quizzes.')
        return redirect('quizzes')

    quiz = get_object_or_404(Quiz, id=quiz_id, is_active=True)

    if quiz.due_date and timezone.now() > quiz.due_date:
        messages.error(request, 'This quiz is closed because the due date has passed.')
        return redirect('quizzes')

    if QuizAttempt.objects.filter(quiz=quiz, student=request.user).exists():
        messages.warning(request, 'You have already attempted this quiz.')
        return redirect('quiz_result', quiz_id=quiz_id)

    questions = quiz.questions.all()

    if not questions.exists():
        messages.error(request, 'This quiz has no questions yet. Please check back later.')
        return redirect('quizzes')

    if request.method == 'POST':
        score = 0
        for question in questions:
            answer = request.POST.get(f'q_{question.id}', '')
            if answer == question.correct_answer:
                score += question.marks

        attempt = QuizAttempt.objects.create(
            quiz=quiz,
            student=request.user,
            score=score,
        )
        _update_engagement(request.user, quiz.subject)
        messages.success(request, f'Quiz submitted! You scored {score}/{quiz.total_marks}')
        return redirect('quiz_result', quiz_id=quiz_id)

    return render(request, 'take_quiz.html', {'quiz': quiz, 'questions': questions})


@login_required
def quiz_result_view(request, quiz_id):
    quiz = get_object_or_404(Quiz, id=quiz_id)
    attempt = get_object_or_404(QuizAttempt, quiz=quiz, student=request.user)
    percentage = (attempt.score / quiz.total_marks * 100) if quiz.total_marks > 0 else 0
    return render(request, 'quiz_result.html', {
        'quiz': quiz,
        'attempt': attempt,
        'percentage': round(percentage, 1),
    })


@login_required
@user_passes_test(is_teacher)
def create_quiz_view(request):
    if request.method == 'POST':
        form = QuizForm(request.POST)
        if form.is_valid():
            quiz = form.save(commit=False)
            quiz.created_by = request.user
            quiz.save()
            messages.success(request, f'Quiz "{quiz.title}" created! Now add questions.')
            return redirect('add_question', quiz_id=quiz.id)
    else:
        form = QuizForm()
    return render(request, 'create_quiz.html', {'form': form})


@login_required
@user_passes_test(is_teacher)
def add_question_view(request, quiz_id):
    quiz = get_object_or_404(Quiz, id=quiz_id)
    if request.method == 'POST':
        form = QuizQuestionForm(request.POST)
        if form.is_valid():
            question = form.save(commit=False)
            question.quiz = quiz
            question.save()
            messages.success(request, 'Question added!')
            if 'add_more' in request.POST:
                return redirect('add_question', quiz_id=quiz_id)
            return redirect('quizzes')
    else:
        form = QuizQuestionForm()
    questions = quiz.questions.all()
    return render(request, 'add_question.html', {'form': form, 'quiz': quiz, 'questions': questions})


# ─────────────────────────────────────────────
# Assignments
# ─────────────────────────────────────────────
@login_required
def assignment_list_view(request):
    if is_teacher(request.user):
        assignments = Assignment.objects.all().order_by('-created_at')
        submissions_pending = AssignmentSubmission.objects.filter(is_graded=False).count()
        return render(request, 'assignments.html', {
            'assignments': assignments,
            'submissions_pending': submissions_pending,
            'is_teacher': True,
        })
    else:
        assignments = Assignment.objects.all().order_by('due_date')
        submitted_ids = set(AssignmentSubmission.objects.filter(
            student=request.user
        ).values_list('assignment_id', flat=True))
        for a in assignments:
            a.submitted = a.id in submitted_ids
        return render(request, 'assignments.html', {
            'assignments': assignments,
            'is_teacher': False,
        })


@login_required
@user_passes_test(is_teacher)
def create_assignment_view(request):
    if request.method == 'POST':
        form = AssignmentForm(request.POST, request.FILES)
        if form.is_valid():
            assignment = form.save(commit=False)
            assignment.created_by = request.user
            assignment.save()
            messages.success(request, 'Assignment created successfully!')
            return redirect('assignments')
    else:
        form = AssignmentForm()
    return render(request, 'create_assignment.html', {'form': form})


@login_required
def submit_assignment_view(request, assignment_id):
    if not is_student(request.user):
        messages.error(request, 'Only students can submit assignments.')
        return redirect('assignments')

    assignment = get_object_or_404(Assignment, id=assignment_id)
    if assignment.due_date and timezone.now() > assignment.due_date:
        messages.error(request, 'Submission closed: the assignment due date has passed.')
        return redirect('assignments')

    if AssignmentSubmission.objects.filter(assignment=assignment, student=request.user).exists():
        messages.warning(request, 'You have already submitted this assignment.')
        return redirect('assignments')
    if request.method == 'POST':
        form = AssignmentSubmissionForm(request.POST, request.FILES)
        if form.is_valid():
            submission = form.save(commit=False)
            submission.assignment = assignment
            submission.student = request.user
            submission.save()
            _update_engagement(request.user, assignment.subject)
            messages.success(request, 'Assignment submitted successfully!')
            return redirect('assignments')
    else:
        form = AssignmentSubmissionForm()
    return render(request, 'submit_assignment.html', {'form': form, 'assignment': assignment})


@login_required
@user_passes_test(is_teacher)
def grade_submission_view(request, submission_id):
    submission = get_object_or_404(AssignmentSubmission, id=submission_id)
    total_marks = submission.assignment.total_marks
    if request.method == 'POST':
        form = GradeSubmissionForm(request.POST, instance=submission, total_marks=total_marks)
        if form.is_valid():
            graded = form.save(commit=False)
            graded.is_graded = True
            graded.save()
            _update_engagement(submission.student, submission.assignment.subject)
            messages.success(request, 'Submission graded successfully!')
            return redirect('assignment_submissions', assignment_id=submission.assignment.id)
    else:
        form = GradeSubmissionForm(instance=submission, total_marks=total_marks)
    return render(request, 'grade_submission.html', {
        'form': form,
        'submission': submission,
        'total_marks': total_marks,
    })


@login_required
@user_passes_test(is_teacher)
def assignment_submissions_view(request, assignment_id):
    assignment = get_object_or_404(Assignment, id=assignment_id)
    submissions = AssignmentSubmission.objects.filter(
        assignment=assignment
    ).select_related('student').order_by('-submitted_at')
    return render(request, 'assignment_submissions.html', {
        'assignment': assignment,
        'submissions': submissions,
    })


# ─────────────────────────────────────────────
# Analytics (Admin / Teacher)
# ─────────────────────────────────────────────
@login_required
@user_passes_test(is_teacher)
def analytics_view(request):
    performances = StudentPerformance.objects.select_related('student', 'subject').all()
    subjects = Subject.objects.all()

    # ── All students for dropdown ──────────────────────────────────────────────
    all_students = User.objects.filter(
        student_profile__isnull=False
    ).order_by('first_name', 'last_name', 'username')

    # ── Filters ────────────────────────────────────────────────────────────────
    subject_id      = request.GET.get('subject', '').strip()
    label_filter    = request.GET.get('label', '').strip()
    risk_only       = request.GET.get('risk_only', '').strip()
    selected_student = request.GET.get('student', '').strip()   # user pk
    sort_by         = request.GET.get('sort', '').strip()

    if subject_id:
        performances = performances.filter(subject_id=subject_id)
    if label_filter:
        performances = performances.filter(predicted_label=label_filter)
    if risk_only == '1':
        performances = performances.filter(is_at_risk=True)
    if selected_student:
        performances = performances.filter(student_id=selected_student)
    if sort_by == 'score_asc':
        performances = performances.order_by('final_exam_score')
    elif sort_by == 'score_desc':
        performances = performances.order_by('-final_exam_score')
    elif sort_by == 'predicted_asc':
        performances = performances.order_by('predicted_score')
    elif sort_by == 'predicted_desc':
        performances = performances.order_by('-predicted_score')
    else:
        performances = performances.order_by('student__username')

    # ── Summary stats ──────────────────────────────────────────────────────────
    avg_score = performances.aggregate(Avg('final_exam_score'))['final_exam_score__avg'] or 0
    at_risk   = performances.filter(is_at_risk=True)

    flipped_avg     = avg_score
    traditional_avg = max(0, avg_score - 12)

    label_counts = {
        'High':    performances.filter(predicted_label='High').count(),
        'Medium':  performances.filter(predicted_label='Medium').count(),
        'Low':     performances.filter(predicted_label='Low').count(),
        'At-Risk': performances.filter(predicted_label='At-Risk').count(),
    }

    return render(request, 'analytics.html', {
        'performances':       performances,
        'subjects':           subjects,
        'all_students':       all_students,
        'selected_subject':   subject_id,
        'selected_label':     label_filter,
        'risk_only':          risk_only,
        'selected_student':   selected_student,
        'sort_by':            sort_by,
        'avg_score':         round(avg_score, 2),
        'at_risk':           at_risk,
        'flipped_avg':       round(flipped_avg, 2),
        'traditional_avg':   round(traditional_avg, 2),
        'label_counts':      json.dumps(label_counts),
        'label_counts_raw':  label_counts,
        'total_records':     performances.count(),
    })


@login_required
@user_passes_test(is_teacher)
def student_detail_analytics_view(request, student_id):
    student = get_object_or_404(User, id=student_id)
    performances = StudentPerformance.objects.filter(student=student).select_related('subject')
    quiz_attempts = QuizAttempt.objects.filter(student=student)
    watch_history = VideoWatchHistory.objects.filter(student=student)

    # Enrich with ML details (same logic as my_performance_view)
    enriched = []
    for p in performances:
        try:
            from ml_model.prediction import predict_student
            pred = predict_student({
                'videos_watched':           p.videos_watched,
                'total_video_time_minutes': p.total_video_time_minutes,
                'quiz_avg_score':           p.quiz_avg_score,
                'assignment_avg_marks':     p.assignment_avg_marks,
                'attendance_percentage':    p.attendance_percentage,
                'participation_score':      p.participation_score,
                'previous_gpa':             p.previous_gpa,
            })
            confidence      = pred['confidence']
            predicted_label = pred['predicted_label']
            is_at_risk      = pred['is_at_risk']
            predicted_score = pred['predicted_score']
        except Exception:
            confidence      = 0
            predicted_label = p.predicted_label or '—'
            is_at_risk      = p.is_at_risk
            predicted_score = p.predicted_score or 0

        features = {
            'Videos':       min(100, int(p.videos_watched * 10)),
            'Video Time':   min(100, int(p.total_video_time_minutes / 3)),
            'Quiz Score':   min(100, int(p.quiz_avg_score * 10)),
            'Assignments':  min(100, int((p.assignment_avg_marks / max(p.assignment_avg_marks, 20)) * 100) if p.assignment_avg_marks else 0),
            'Attendance':   min(100, int(p.attendance_percentage)),
            'Participation':min(100, int(p.participation_score * 10)),
        }
        weakest = min(features, key=features.get)
        enriched.append({
            'record':           p,
            'confidence':       confidence,
            'predicted_label':  predicted_label,
            'is_at_risk':       is_at_risk,
            'predicted_score':  predicted_score,
            'features':         features,
            'feature_keys':     list(features.keys()),
            'feature_vals':     list(features.values()),
            'weakest':          weakest,
            'recommendation':   _get_recommendation(weakest),
        })

    return render(request, 'student_analytics.html', {
        'student':              student,
        'enriched':             enriched,
        'performances':         performances,
        'quiz_attempts':        quiz_attempts,
        'watch_history':        watch_history,
        'total_videos_watched': watch_history.filter(completed=True).count(),
        'avg_quiz':             quiz_attempts.aggregate(Avg('score'))['score__avg'] or 0,
    })


# ─────────────────────────────────────────────
# ML Prediction
# ─────────────────────────────────────────────
@login_required
@user_passes_test(is_teacher)
def run_ml_prediction_view(request):
    if request.method != 'POST':
        messages.error(request, 'ML prediction must be triggered via POST.')
        return redirect('analytics')
    from ml_model.prediction import predict_all_students
    try:
        results = predict_all_students()
        messages.success(request, f'ML prediction complete! Processed {len(results)} records.')
    except Exception as e:
        logger.exception('run_ml_prediction_view failed: %s', e)
        messages.error(request, 'ML prediction failed. Please try again later or check server logs.')
    return redirect('analytics')


# ─────────────────────────────────────────────
# ML Recommendation Helper
# ─────────────────────────────────────────────

def _get_recommendation(weakest: str) -> str:
    """Return a human-readable recommendation based on the student's weakest metric."""
    return {
        'Videos':        "📹 Watch more video lectures — aim for at least 5 videos per chapter.",
        'Video Time':    "⏱️ Spend more time on each video — don't skip; re-watch difficult sections.",
        'Quiz Score':    "📝 Practice more quizzes — review wrong answers and revisit that chapter.",
        'Assignments':   "📋 Submit assignments on time and re-read the grading feedback.",
        'Attendance':    "📅 Improve your class attendance — being present dramatically raises your score.",
        'Participation': "💬 Ask more questions in the AI Tutor — each question builds your participation score.",
    }.get(weakest, "📚 Keep revising regularly and stay consistent across all subjects.")


@login_required
def my_performance_view(request):
    performance_records = StudentPerformance.objects.filter(
        student=request.user
    ).select_related('subject')
    quiz_attempts   = QuizAttempt.objects.filter(student=request.user)
    watch_history   = VideoWatchHistory.objects.filter(student=request.user, completed=True)
    submissions     = AssignmentSubmission.objects.filter(student=request.user, is_graded=True)

    # ── Enrich each record with live ML prediction details ─────────────────────
    enriched = []
    for p in performance_records:
        try:
            from ml_model.prediction import predict_student
            pred = predict_student({
                'videos_watched':           p.videos_watched,
                'total_video_time_minutes': p.total_video_time_minutes,
                'quiz_avg_score':           p.quiz_avg_score,
                'assignment_avg_marks':     p.assignment_avg_marks,
                'attendance_percentage':    p.attendance_percentage,
                'participation_score':      p.participation_score,
                'previous_gpa':             p.previous_gpa,
            })
            confidence      = pred['confidence']
            predicted_label = pred['predicted_label']
            is_at_risk      = pred['is_at_risk']
            predicted_score = pred['predicted_score']
        except Exception:
            confidence      = 0
            predicted_label = p.predicted_label or '—'
            is_at_risk      = p.is_at_risk
            predicted_score = p.predicted_score or 0

        # Normalize each feature to 0–100 for progress bars & radar chart
        features = {
            'Videos':       min(100, int(p.videos_watched * 10)),
            'Video Time':   min(100, int(p.total_video_time_minutes / 3)),
            'Quiz Score':   min(100, int(p.quiz_avg_score * 10)),
            'Assignments':  min(100, int((p.assignment_avg_marks / max(p.assignment_avg_marks, 20)) * 100) if p.assignment_avg_marks else 0),
            'Attendance':   min(100, int(p.attendance_percentage)),
            'Participation':min(100, int(p.participation_score * 10)),
        }
        weakest = min(features, key=features.get)
        enriched.append({
            'record':           p,
            'confidence':       confidence,
            'predicted_label':  predicted_label,
            'is_at_risk':       is_at_risk,
            'predicted_score':  predicted_score,
            'features':         features,
            'feature_keys':     list(features.keys()),
            'feature_vals':     list(features.values()),
            'weakest':          weakest,
            'recommendation':   _get_recommendation(weakest),
        })

    return render(request, 'my_performance.html', {
        'enriched':       enriched,
        'quiz_attempts':  quiz_attempts,
        'watch_history':  watch_history,
        'submissions':    submissions,
        'avg_quiz':       quiz_attempts.aggregate(Avg('score'))['score__avg'] or 0,
        'total_watched':  watch_history.count(),
    })


# ─────────────────────────────────────────────
# Export CSV
# ─────────────────────────────────────────────
@login_required
@user_passes_test(is_teacher)
def export_performance_csv(request):
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="student_performance.csv"'
    writer = csv.writer(response)
    writer.writerow([
        'Student', 'Roll No', 'Subject', 'Videos Watched', 'Total Video Time (min)',
        'Quiz Avg Score', 'Assignment Avg', 'Attendance %', 'Participation',
        'Previous GPA', 'Final Score', 'Performance Label', 'Predicted Score',
        'Predicted Label', 'At Risk'
    ])
    for perf in StudentPerformance.objects.select_related('student', 'subject').all():
        roll = ''
        try:
            roll = perf.student.student_profile.roll_number
        except Exception:
            pass
        writer.writerow([
            perf.student.get_full_name(),
            roll,
            perf.subject.name,
            perf.videos_watched,
            perf.total_video_time_minutes,
            perf.quiz_avg_score,
            perf.assignment_avg_marks,
            perf.attendance_percentage,
            perf.participation_score,
            perf.previous_gpa,
            perf.final_exam_score,
            perf.performance_label,
            perf.predicted_score or '',
            perf.predicted_label or '',
            perf.is_at_risk,
        ])
    return response


# ─────────────────────────────────────────────
# Notifications
# ─────────────────────────────────────────────
@login_required
def mark_notification_read(request, notif_id):
    notif = get_object_or_404(Notification, id=notif_id, recipient=request.user)
    notif.is_read = True
    notif.save()
    referer = request.META.get('HTTP_REFERER', '')
    if referer and url_has_allowed_host_and_scheme(
        referer,
        allowed_hosts={request.get_host()},
        require_https=request.is_secure(),
    ):
        return redirect(referer)
    return redirect('dashboard')


# ─────────────────────────────────────────────
# Internal helper
# ─────────────────────────────────────────────
def _update_engagement(user, subject):
    """Recalculate and update StudentPerformance aggregates for a user/subject,
    then immediately run ML prediction and save predicted fields."""
    perf, _ = StudentPerformance.objects.get_or_create(student=user, subject=subject)

    # ── Videos ────────────────────────────────────────────────────────────────
    videos_watched = VideoWatchHistory.objects.filter(
        student=user, video__subject=subject, completed=True
    )
    perf.videos_watched = videos_watched.count()
    perf.total_video_time_minutes = videos_watched.aggregate(
        total=Coalesce(Sum('watch_duration_minutes'), 0.0, output_field=FloatField())
    )['total']

    # ── Quizzes ───────────────────────────────────────────────────────────────
    quiz_attempts = QuizAttempt.objects.filter(student=user, quiz__subject=subject)
    if quiz_attempts.exists():
        perf.quiz_avg_score = quiz_attempts.aggregate(Avg('score'))['score__avg'] or 0

    # ── Assignments ────────────────────────────────────────────────────────────
    submissions = AssignmentSubmission.objects.filter(
        student=user, assignment__subject=subject, is_graded=True
    )
    if submissions.exists():
        perf.assignment_avg_marks = submissions.aggregate(Avg('marks_obtained'))['marks_obtained__avg'] or 0

    # ── Attendance ─────────────────────────────────────────────────────────────
    attendance = Attendance.objects.filter(student=user, subject=subject)
    if attendance.count() > 0:
        present_count = attendance.filter(present=True).count()
        perf.attendance_percentage = (present_count / attendance.count()) * 100

    # ── Chat Participation (0.5 pts per question asked, capped at 10) ──────────
    chat_count = ChatMessage.objects.filter(
        student=user, subject=subject, role='user'
    ).count()
    perf.participation_score = min(10.0, round(chat_count * 0.5, 1))

    # ── Previous GPA ───────────────────────────────────────────────────────────
    try:
        perf.previous_gpa = user.student_profile.previous_gpa
    except Exception:
        pass

    perf.save()

    # ── Auto ML Prediction ─────────────────────────────────────────────────────
    try:
        from ml_model.prediction import predict_student
        pred = predict_student({
            'videos_watched':           perf.videos_watched,
            'total_video_time_minutes': perf.total_video_time_minutes,
            'quiz_avg_score':           perf.quiz_avg_score,
            'assignment_avg_marks':     perf.assignment_avg_marks,
            'attendance_percentage':    perf.attendance_percentage,
            'participation_score':      perf.participation_score,
            'previous_gpa':             perf.previous_gpa,
        })
        perf.predicted_score = pred['predicted_score']
        perf.predicted_label = pred['predicted_label']
        perf.is_at_risk       = pred['is_at_risk']
        perf.save(update_fields=['predicted_score', 'predicted_label', 'is_at_risk'])

        # Send at-risk notification (once per subject, only if not already sent)
        if pred['is_at_risk']:
            already = Notification.objects.filter(
                recipient=user, is_read=False,
                message__icontains=subject.name,
            ).exists()
            if not already:
                Notification.objects.create(
                    recipient=user,
                    message=(
                        f"⚠️ Your performance in {subject.name} needs attention. "
                        f"ML Predicted Score: {pred['predicted_score']} ({pred['predicted_label']}). "
                        "Try watching more videos, asking questions, and taking quizzes."
                    ),
                )
    except Exception:
        pass  # Models not yet trained — silently skip


# ─────────────────────────────────────────────
# RAG Chatbot Views
# ─────────────────────────────────────────────

@login_required
def chat_ask_view(request):
    """Handle a student's chat question. POST only. Returns JSON."""
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)

    try:
        data = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)

    user_query = data.get('query', '').strip()
    subject_code = data.get('subject_code', '').strip() or None
    lang_code = data.get('lang_code', '').strip() or None

    if not user_query:
        return JsonResponse({'error': 'Query is empty'}, status=400)

    # Build chat history for context (last 6 messages)
    recent_msgs = (
        ChatMessage.objects
        .filter(student=request.user)
        .order_by('-created_at')[:6]
    )
    chat_history = [
        {'role': m.role, 'content': m.content}
        for m in reversed(recent_msgs)
    ]

    # Call RAG engine
    try:
        from rag_engine.chat import ask
        result = ask(
            user_query=user_query,
            subject_code=subject_code,
            chat_history=chat_history,
            top_k=5,
            lang_pref=lang_code,
        )
    except Exception as e:
        logger.exception('chat_ask_view failed: %s', e)
        result = {
            'reply': 'Sorry, something went wrong while generating the response.',
            'sources': [],
            'error': 'internal_error',
        }

    reply = result.get('reply', '')
    sources = result.get('sources', [])

    # Persist messages
    subject_obj = None
    if subject_code:
        try:
            subject_obj = Subject.objects.get(code=subject_code)
        except Subject.DoesNotExist:
            pass

    ChatMessage.objects.create(
        student=request.user,
        subject=subject_obj,
        role='user',
        content=user_query,
    )
    ChatMessage.objects.create(
        student=request.user,
        subject=subject_obj,
        role='assistant',
        content=reply,
        sources=", ".join(sources),
    )

    # Update participation score from chat activity
    if subject_obj and is_student(request.user):
        try:
            _update_engagement(request.user, subject_obj)
        except Exception:
            pass

    return JsonResponse({'reply': reply, 'sources': sources})


@login_required
def chat_history_view(request):
    """Return last 20 chat messages for the current user as JSON."""
    messages_qs = (
        ChatMessage.objects
        .filter(student=request.user)
        .order_by('-created_at')[:20]
    )
    data = [
        {
            'role': m.role,
            'content': m.content,
            'subject': m.subject.code if m.subject else '',
            'sources': m.sources,
            'created_at': m.created_at.strftime('%H:%M'),
        }
        for m in reversed(messages_qs)
    ]
    return JsonResponse({'messages': data})


@login_required
def chat_stream_view(request):
    """SSE streaming view — yields Perplexity-like events."""
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)

    try:
        data = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)

    user_query = data.get('query', '').strip()
    subject_code = data.get('subject_code', '').strip() or None
    lang_code = data.get('lang_code', '').strip() or None

    if not user_query:
        return JsonResponse({'error': 'Empty query'}, status=400)

    # Build recent chat history
    recent_msgs = (
        ChatMessage.objects
        .filter(student=request.user)
        .order_by('-created_at')[:6]
    )
    chat_history = [
        {'role': m.role, 'content': m.content}
        for m in reversed(recent_msgs)
    ]

    def event_stream():
        from rag_engine.chat import stream_answer
        full_reply = ''
        sources_list = []
        try:
            for event_str in stream_answer(
                user_query=user_query,
                subject_code=subject_code,
                chat_history=chat_history,
                top_k=3,
                lang_pref=lang_code,
            ):
                # event_str is already "data: {...}\n\n"
                yield event_str
                # capture full_reply and sources from done event
                if event_str.startswith('data: '):
                    try:
                        import json as _json
                        payload = _json.loads(event_str[6:])
                        if payload.get('type') == 'done':
                            full_reply = payload.get('full_reply', '')
                            sources_list = payload.get('sources', [])
                    except Exception:
                        pass
        except Exception as e:
            logger.exception('chat_stream_view failed: %s', e)
            import json as _json
            yield f'data: {_json.dumps({"type": "error", "text": "Sorry, the chatbot is temporarily unavailable."})}\n\n'
            return

        # Persist to DB after stream completes
        try:
            subject_obj = None
            if subject_code:
                try:
                    subject_obj = Subject.objects.get(code=subject_code)
                except Subject.DoesNotExist:
                    pass
            ChatMessage.objects.create(
                student=request.user,
                subject=subject_obj,
                role='user',
                content=user_query,
            )
            ChatMessage.objects.create(
                student=request.user,
                subject=subject_obj,
                role='assistant',
                content=full_reply,
                sources=', '.join(sources_list),
            )
            # Update participation score from chat activity
            if subject_obj and is_student(request.user):
                try:
                    _update_engagement(request.user, subject_obj)
                except Exception:
                    pass
        except Exception:
            pass

    response = StreamingHttpResponse(
        event_stream(),
        content_type='text/event-stream',
    )
    response['Cache-Control'] = 'no-cache'
    response['X-Accel-Buffering'] = 'no'
    return response


@login_required
def chat_pdf_view(request):
    """Accept a PDF or Word (.docx) upload + optional question, stream an explanation via Groq SSE."""
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)

    uploaded_file = request.FILES.get('pdf_file')
    if not uploaded_file:
        return JsonResponse({'error': 'No file uploaded'}, status=400)

    fname = uploaded_file.name.lower()
    is_pdf  = fname.endswith('.pdf')
    is_docx = fname.endswith('.docx')

    if not is_pdf and not is_docx:
        return JsonResponse({'error': 'Only PDF and Word (.docx) files are supported'}, status=400)

    if uploaded_file.size > 10 * 1024 * 1024:   # 10 MB limit
        return JsonResponse({'error': 'File too large (max 10 MB)'}, status=400)

    user_query = request.POST.get('query', '').strip()

    # Extract text — branch by file type
    try:
        import io
        raw_bytes = uploaded_file.read()
        if is_pdf:
            import pdfplumber, logging
            logging.getLogger('pdfminer').setLevel(logging.ERROR)
            with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
                pages_text = []
                for page in pdf.pages[:20]:   # max 20 pages
                    t = page.extract_text()
                    if t:
                        pages_text.append(t.strip())
            pdf_text = '\n\n'.join(pages_text)
        else:  # .docx
            from docx import Document
            doc = Document(io.BytesIO(raw_bytes))
            paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            # Also grab table cell text
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(c.text.strip() for c in row.cells if c.text.strip())
                    if row_text:
                        paras.append(row_text)
            pdf_text = '\n\n'.join(paras)
    except Exception as e:
        logger.exception('chat_pdf_view failed to parse uploaded file: %s', e)
        return JsonResponse({'error': 'Could not read this file. Please upload a valid PDF or .docx file.'}, status=400)

    if not pdf_text.strip():
        return JsonResponse({'error': 'File appears to be empty or image-only (no extractable text)'}, status=400)

    # Truncate to avoid large context
    MAX_PDF_CHARS = 6000
    truncated = pdf_text[:MAX_PDF_CHARS]
    was_truncated = len(pdf_text) > MAX_PDF_CHARS

    pdf_name = uploaded_file.name

    def event_stream():
        import json as _json
        from groq import Groq
        from django.conf import settings as _settings

        def sse(payload):
            return f'data: {_json.dumps(payload)}\n\n'

        # Yield a synthetic "source" card for the uploaded file
        file_type_label = 'PDF' if is_pdf else 'Word'
        yield sse({'type': 'sources', 'sources': [{'title': pdf_name, 'subject': file_type_label}]})

        system_prompt = (
            'You are FlipLearn AI, an expert academic tutor for M.Tech CSE students. '
            f'A student has uploaded a {file_type_label} document. Analyse it carefully and respond helpfully. '
            'Use **bold** for key terms, `code` for code/commands, and ### for section headers. '
            'Be thorough but concise.'
        )

        if user_query:
            user_content = (
                f'{file_type_label} content ({pdf_name})'
                + (' [truncated to first 6000 chars]' if was_truncated else '') +
                f':\n\n{truncated}\n\n---\n\nStudent question: {user_query}'
            )
        else:
            user_content = (
                f'{file_type_label} content ({pdf_name})'
                + (' [truncated to first 6000 chars]' if was_truncated else '') +
                f':\n\n{truncated}\n\n---\n\n'
                'Please provide a comprehensive explanation of this document: '
                'cover the main topics, key concepts, important points, and any conclusions.'
            )

        messages = [
            {'role': 'system', 'content': system_prompt},
            {'role': 'user',   'content': user_content},
        ]

        full_reply = ''
        try:
            api_key = getattr(_settings, 'GROQ_API_KEY', '')
            client = Groq(api_key=api_key)
            stream = client.chat.completions.create(
                model='llama-3.1-8b-instant',
                messages=messages,
                max_tokens=1200,
                temperature=0.3,
                stream=True,
            )
            for chunk in stream:
                token = chunk.choices[0].delta.content or ''
                if token:
                    full_reply += token
                    yield sse({'type': 'token', 'text': token})
        except Exception as e:
            logger.exception('chat_pdf_view stream failed: %s', e)
            yield sse({'type': 'error', 'text': 'Sorry, the AI service is temporarily unavailable.'})

        yield sse({'type': 'done', 'full_reply': full_reply, 'sources': [pdf_name]})

        # Persist to DB
        try:
            label = user_query if user_query else f'Explain PDF: {pdf_name}'
            ChatMessage.objects.create(student=request.user, role='user',    content=label)
            ChatMessage.objects.create(student=request.user, role='assistant', content=full_reply,
                                       sources=pdf_name)
        except Exception:
            pass

    resp = StreamingHttpResponse(event_stream(), content_type='text/event-stream')
    resp['Cache-Control'] = 'no-cache'
    resp['X-Accel-Buffering'] = 'no'
    return resp


# ─────────────────────────────────────────────
# RAG Index Rebuild (Admin only)
# ─────────────────────────────────────────────
@login_required
def rebuild_rag_view(request):
    """AJAX endpoint – rebuilds the FAISS RAG index. Staff/admin only."""
    if not request.user.is_superuser:
        return JsonResponse({'status': 'error', 'message': 'Permission denied.'}, status=403)
    if request.method == 'GET':
        with RAG_REBUILD_LOCK:
            return JsonResponse({'status': 'ok', **_rag_status_payload()})

    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'message': 'POST required.'}, status=405)

    with RAG_REBUILD_LOCK:
        if RAG_REBUILD_STATE['running']:
            return JsonResponse({
                'status': 'ok',
                'message': 'RAG rebuild is already running.',
                **_rag_status_payload(),
            })

        RAG_REBUILD_STATE['running'] = True
        RAG_REBUILD_STATE['started_at'] = timezone.now()
        RAG_REBUILD_STATE['finished_at'] = None
        RAG_REBUILD_STATE['last_error'] = None

    if getattr(settings, 'RAG_REBUILD_SYNC', False):
        _run_rag_rebuild_job()
        with RAG_REBUILD_LOCK:
            payload = _rag_status_payload()
        if payload['last_error']:
            return JsonResponse({'status': 'error', 'message': payload['last_error'], **payload}, status=500)
        return JsonResponse({'status': 'ok', 'message': 'RAG index rebuilt successfully.', **payload})

    threading.Thread(target=_run_rag_rebuild_job, daemon=True, name='fliplearn-rag-rebuild').start()
    with RAG_REBUILD_LOCK:
        payload = _rag_status_payload()
    return JsonResponse({
        'status': 'ok',
        'message': 'RAG rebuild started in background. It may take several minutes.',
        **payload,
    })
